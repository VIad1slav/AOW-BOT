#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""AOW Document Converter  — Pro Forma → Invoice / Spec cleanup"""

import threading, os, re, json, copy
from pathlib import Path
from datetime import datetime
from decimal import Decimal, InvalidOperation
from collections import Counter, defaultdict

# ─── ANTI-CIRCUMVENTION CLAUSE ───────────────────────────────────────────────

# Pre-compiled regex patterns (faster than re.compile on each call)
_RE_DATE      = re.compile(r'\d{4}-\d{2}-\d{2}')
_RE_CELL_REF  = re.compile(r'(\$?)([A-Z]+)(\$?)(\d+)')
_RE_EUR_AMT   = re.compile(r'[\d\u00a0 ]+,\d{2}')
_RE_FOOTER    = re.compile(r'\b109\d{4}')

ANTI_CIRCUMVENTION = (
    "(1) [Importer/Buyer] shall not sell, export or re-export - either directly or indirectly - "
    "to the Russian Federation or for use in the Russian Federation; or to Belarus or for use in Belarus "
    "goods supplied under or in connection with this Agreement. The Agreement shall apply to:\n"
    "- goods identified in Article 12g of Council Regulation (EU) No 833/2014; and\n"
    "- goods identified in Article 8g of Council Regulation (EC) No 765/2006.\n"
    "(2) The [Importer/Buyer] shall use its best endeavours to ensure that the purpose of paragraph (1) "
    "is not frustrated by third parties further down the trade chain, including any resellers.\n"
    "(3) The [Importer/Buyer] shall establish and maintain an appropriate control mechanism to detect "
    "conduct by third parties down the supply chain that would frustrate the purpose of paragraph 1.\n"
    "(4) Any violation of paragraphs (1), (2) or (3) shall constitute a material breach of an essential "
    "element of this Agreement and [Exporter] shall have the right to appropriate remedies, including:\n"
    "- termination of this Agreement; and\n"
    "- an appropriate penalty of 10% of the total value of this Agreement or of the price of the "
    "exported goods, whichever is greater.\n"
    "(5) The [Importer/Buyer] shall promptly inform the [Exporter/Seller] of any problems with the "
    "application of paragraphs (1), (2) or (3), including any relevant activities of third parties "
    "that may frustrate the purpose of paragraph (1).\n"
    "The [Importer/Buyer] shall provide the [Exporter/Shipper] with information relevant to the "
    "fulfilment of the obligations under paragraphs 1, 2 and 3 within two weeks of the simple request "
    "for such information."
)

# ─── DOCX HELPERS ────────────────────────────────────────────────────────────

def _txt(para):
    return ''.join(r.text for r in para.runs)

def _set(para, text):
    """Set paragraph text, preserving first run's formatting."""
    if para.runs:
        para.runs[0].text = text
        for r in para.runs[1:]:
            r.text = ''
    else:
        para.add_run(text)

def get_articles_from_xlsx(xlsx_path, log):
    """Return set of article strings from XLSX spec."""
    import openpyxl
    wb = openpyxl.load_workbook(xlsx_path, data_only=True)
    ws = wb.active
    art_col, hdr_row = None, None
    for r in range(1, 9):
        for c in range(1, 5):
            v = ws.cell(r, c).value
            if v and any(k in str(v) for k in ('Artikel', 'Article', 'Артикул')):
                art_col, hdr_row = c, r; break
        if art_col: break
    if art_col is None:
        log('  ⚠ Колонка Artikel не найдена в XLSX')
        return set()
    arts = set()
    for r in range(hdr_row + 1, ws.max_row + 1):
        v = ws.cell(r, art_col).value
        if v:
            s = str(v).strip()
            if s and not any(w in s for w in ('Итого', 'Total', 'ИТОГО', 'TOTAL')):
                arts.add(s)
    log(f'  ✓ Найдено {len(arts)} артикулов в спецификации')
    return arts

def get_xlsx_items_data(xlsx_path, log):
    """Return all item rows from XLSX as list of dicts."""
    import openpyxl
    wb = openpyxl.load_workbook(xlsx_path, data_only=True)
    ws = wb.active
    art_col, hdr_row = None, None
    for r in range(1, 9):
        for c in range(1, 5):
            v = ws.cell(r, c).value
            if v and any(k in str(v) for k in ('Artikel', 'Article', 'Артикул')):
                art_col, hdr_row = c, r; break
        if art_col: break
    if art_col is None: return []
    # Detect English description column (col after article)
    desc_col = art_col + 2   # col 4 = English description in original XLSX
    qty_col  = art_col + 7   # col 9  = Quantity pcs
    price_col= art_col + 10  # col 12 = E-Price
    total_col= art_col + 11  # col 13 = Total EUR
    items = []
    for r in range(hdr_row+1, ws.max_row+1):
        art = str(ws.cell(r, art_col).value or '').strip()
        qty = ws.cell(r, qty_col).value
        if not art or not qty: continue
        if any(w in art for w in ('Итого','Total','ИТОГО','TOTAL','Условия','Delivery','Payment','Shipment')): continue
        items.append({
            'article': art,
            'desc':    str(ws.cell(r, desc_col).value or '').strip(),
            'qty':     qty,
            'price':   ws.cell(r, price_col).value,
            'total':   ws.cell(r, total_col).value,
        })
    log(f'  ✓ XLSX: {len(items)} строк данных')
    return items

# ─── DOCX MAIN ────────────────────────────────────────────────────────────────

def process_docx(src, dst, params, spec_arts, log):
    from docx import Document
    from decimal import Decimal, InvalidOperation

    inv_num  = params['invoice_num']
    date_str = params['date']
    pf_ref   = params['pf_ref']

    try:
        d, m, y = date_str.split('.')
        date_word = f'{y}-{m.zfill(2)}-{d.zfill(2)}'
    except Exception:
        date_word = date_str

    log('📄 Обрабатываю DOCX...')
    doc = Document(src)

    # ── 1. TITLE ──────────────────────────────────────────────────────────────
    for p in doc.paragraphs:
        t = _txt(p)
        if 'Faktura Pro Forma' in t and 'invoice' in t.lower():
            _set(p, f'Faktura / Invoice {inv_num}')
            log(f'  ✓ Заголовок → «Faktura / Invoice {inv_num}»')
            break

    # ── 2. DATE — keep label bold, set value non-bold ─────────────────────────
    for p in doc.paragraphs:
        t = _txt(p)
        if 'Data wystawienia' in t and re.search(r'\d{4}-\d{2}-\d{2}', t):
            # Find the run containing 'Warszawa' and set date there (non-bold)
            for i, run in enumerate(p.runs):
                if 'Warszawa' in run.text:
                    run.text = f'Warszawa, {date_word}'
                    run.bold = False
                    run.font.size = p.runs[0].font.size
                    for r in p.runs[i+1:]:
                        r.text = ''
                        r.bold = False
                    break
            log(f'  ✓ Дата → {date_word}')
            break

    # ── 3. DELETE "Termin płatności" paragraph ────────────────────────────────
    for p in doc.paragraphs:
        if 'Termin p' in _txt(p):   # matches 'Termin płatności'
            p._element.getparent().remove(p._element)
            log('  ✓ Удалена строка Termin płatności')
            break

    # ── 4. PAYMENT — remove non-bold Transfer, add bold PF reference ───────────
    for p in doc.paragraphs:
        t = _txt(p)
        if 'Przelew / Transfer' in t and pf_ref not in t:
            _ref = p.runs[0]
            # Remove 'Transfer' from non-bold run, keep ' / '
            for r in p.runs:
                if r.text == ' / Transfer':
                    r.text = ' / '
                elif r.text.strip() == '' and not r.bold:
                    r.text = ''
            # Add bold reference without leading space
            run = p.add_run(f'Transfer Faktura Pro Forma / Pro Forma {pf_ref}')
            run.bold = True
            run.font.size = _ref.font.size
            run.font.name = _ref.font.name
            log(f'  ✓ Добавлена ссылка (bold): {pf_ref}')
            break

    # ── 5. ITEM TABLE — sync with XLSX (remove extra, add missing) ─────────────
    table = doc.tables[0]
    removed_items = []
    added_items = []

    # 5a. Remove rows from DOCX not in XLSX
    if spec_arts:
        to_del = []
        for row in table.rows[1:]:
            art = _txt(row.cells[1].paragraphs[0]).strip() if row.cells[1].paragraphs else ''
            if art and 'Razem' not in art and 'Total' not in art:
                if art not in spec_arts:
                    to_del.append((row, art))
        for row, art in to_del:
            row._tr.getparent().remove(row._tr)
            removed_items.append(art)
            log(f'    ↳ Удалена позиция: {art}')
        if removed_items:
            log(f'  ✓ Удалено позиций: {len(removed_items)}')

    # 5b. Add rows from XLSX missing in DOCX
    xlsx_items_all = params.get('_xlsx_items', [])
    if spec_arts and xlsx_items_all:
        # Build dict of existing DOCX quantities per article (for matching)
        pf_qtys = {}
        for row in table.rows[1:]:
            art = _txt(row.cells[1].paragraphs[0]).strip() if row.cells[1].paragraphs else ''
            row_full = ''.join(_txt(p) for cell in row.cells for p in cell.paragraphs)
            qty_s = _txt(row.cells[3].paragraphs[0]).replace('\u00a0','').replace(' ','') if row.cells[3].paragraphs else '0'
            if art and 'Razem' not in row_full:
                try:
                    if art not in pf_qtys: pf_qtys[art] = []
                    pf_qtys[art].append(int(float(qty_s)))
                except: pass

        # Match each XLSX row against existing DOCX rows by quantity
        # Unmatched XLSX rows need to be added
        remaining_pf = {art: list(qtys) for art, qtys in pf_qtys.items()}
        to_add = []
        for item in xlsx_items_all:
            art = item['article']
            try: qty = int(item['qty'])
            except: qty = 0
            if art in remaining_pf and qty in remaining_pf[art]:
                remaining_pf[art].remove(qty)  # matched — skip
            else:
                to_add.append(item)             # unmatched — add

        if to_add:
            def _fmt(v):
                try: return f'{float(v):,.2f}'.replace(',', ' ').replace('.', ',')
                except: return str(v or 0)

            razem_tr   = table.rows[-1]._tr
            template_tr = table.rows[1]._tr

            for item in to_add:
                new_tr = copy.deepcopy(template_tr)
                razem_tr.addprevious(new_tr)
                for row in table.rows:
                    if row._tr is new_tr:
                        vals = ['0', item['article'], item['desc'],
                                str(int(item['qty'])), _fmt(item['price']),
                                _fmt(item['total']), '0', '0', _fmt(item['total'])]
                        for ci, val in enumerate(vals):
                            if ci < len(row.cells) and row.cells[ci].paragraphs:
                                _set(row.cells[ci].paragraphs[0], val)
                        break
                added_items.append(item['article'])
                log(f"    ↳ Добавлена позиция: {item['article']} ({int(item['qty'])} шт)")
            log(f'  ✓ Добавлено позиций: {len(added_items)}')

    # 5c. Reorder DOCX rows to match XLSX order
    if xlsx_items_all:
        from collections import defaultdict
        razem_tr   = table.rows[-1]._tr
        header_tr  = table.rows[0]._tr

        # Index current rows by (article, qty)
        rows_by_key = defaultdict(list)
        for row in table.rows[1:]:
            row_full = ''.join(_txt(p) for cell in row.cells for p in cell.paragraphs)
            if 'Razem' in row_full: continue
            art   = _txt(row.cells[1].paragraphs[0]).strip() if row.cells[1].paragraphs else ''
            qty_s = _txt(row.cells[3].paragraphs[0]).replace('\u00a0','').replace(' ','') if row.cells[3].paragraphs else '0'
            try: qty = int(float(qty_s))
            except: qty = 0
            rows_by_key[(art, qty)].append(row._tr)

        # Detach all data rows
        for tr_el in list(rows_by_key[k] for k in rows_by_key for _ in [0]):
            pass  # just iterating — actual detach below
        all_data_trs = [tr for key, trs in rows_by_key.items() for tr in trs]
        for tr in all_data_trs:
            tr.getparent().remove(tr)

        # Re-insert in XLSX order
        used = defaultdict(int)
        inserted = 0
        for item in xlsx_items_all:
            art = item['article']
            try: qty = int(item['qty'])
            except: qty = 0
            key = (art, qty)
            if key in rows_by_key and used[key] < len(rows_by_key[key]):
                tr = rows_by_key[key][used[key]]
                razem_tr.addprevious(tr)
                used[key] += 1
                inserted += 1

        log(f'  ✓ Порядок строк → как в XLSX ({inserted} позиций)')

    # 5d. Renumber all items sequentially
    lp = 1
    for row in table.rows[1:]:
        row_full = ''.join(_txt(p) for cell in row.cells for p in cell.paragraphs)
        if 'Razem' in row_full: break
        if row.cells[0].paragraphs:
            _set(row.cells[0].paragraphs[0], str(lp)); lp += 1
    log(f'  ✓ Перенумеровано позиций: {lp - 1}')

    # ── 6. RECALCULATE TOTALS ─────────────────────────────────────────────────
    total_qty = 0
    total_net = Decimal('0')
    razem_row = None

    for row in table.rows[1:]:
        # Razem row: text «Razem / Total» is in Cell[2], not Cell[1] - check all cells
        row_full = ''.join(_txt(p) for cell in row.cells for p in cell.paragraphs)
        if 'Razem' in row_full:
            razem_row = row
            continue
        art_t = _txt(row.cells[1].paragraphs[0]).strip() if row.cells[1].paragraphs else ''
        qty_s = _txt(row.cells[3].paragraphs[0]).replace('\u00a0','').replace(' ','').replace(',','.') if row.cells[3].paragraphs else ''
        net_s = _txt(row.cells[5].paragraphs[0]).replace('\u00a0','').replace(' ','').replace(',','.') if row.cells[5].paragraphs else ''
        try: total_qty += int(float(qty_s))
        except: pass
        try: total_net += Decimal(net_s)
        except: pass

    def fmt_qty(n):   return f'{n:,}'.replace(',', ' ')
    def fmt_eur(n):   return f'{float(n):,.2f}'.replace(',', ' ').replace('.', ',')

    qty_str = fmt_qty(total_qty)
    eur_str = fmt_eur(total_net)

    if razem_row and total_qty > 0:
        for ci, val in [(3, qty_str), (5, eur_str), (8, eur_str)]:
            if razem_row.cells[ci].paragraphs:
                _set(razem_row.cells[ci].paragraphs[0], val)
        log(f'  ✓ Итого: {qty_str} шт., {eur_str} EUR')

        # Update standalone total paragraphs (Wartość netto / brutto)
        for p in doc.paragraphs:
            t = _txt(p)
            if ('Warto' in t or 'Wartość' in t) and 'EUR' in t and 'VAT' not in t:
                for run in p.runs:
                    if re.search(r'[\d\u00a0 ]+,\d{2}', run.text):
                        run.text = re.sub(r'[\d\u00a0 ]+,\d{2}', eur_str.replace(' ', '\u00a0'), run.text)

    # ── 7. UWAGI / ATTENTION ──────────────────────────────────────────────────
    for p in doc.paragraphs:
        if 'Uwagi' in _txt(p):
            # Rebuild runs with correct bold formatting + preserve font
            _sz = p.runs[0].font.size
            _nm = p.runs[0].font.name
            p.runs[0].text = 'Uwagi / Attention:\nWarunki dostawy / Delivery terms: '
            p.runs[0].bold = True
            for r in p.runs[1:]: r.text = ''
            def _ar(txt, bold):
                r = p.add_run(txt); r.bold = bold
                r.font.size = _sz; r.font.name = _nm; return r
            _ar('FCA Warszaw Poland \n', False)
            _ar('Wewnątrzwspólnotowa dostawa towarów (WDT)\n', True)
            _ar(ANTI_CIRCUMVENTION, False)
            log('  ✓ Uwagi / Attention обновлено (bold структура)')
            break

    # ── 8. CLEAR FOOTER REFERENCE ─────────────────────────────────────────────
    cleared = False
    for p in doc.paragraphs:
        if re.search(r'\b10\d{5}\b', _txt(p)):
            for run in p.runs: run.text = ''
            cleared = True
    for p in list(doc.paragraphs):
        t = _txt(p).strip()
        if t and re.fullmatch(r'[\d\s\u00a0]+', t):
            p._element.getparent().remove(p._element)
            cleared = True
    if cleared:
        log('  ✓ Сноска / хвостовые номера очищены')

    # ── Fix price format: 4 decimal → 2 decimal in Cell[4] ─────────────────────
    for row in doc.tables[0].rows[1:]:
        row_full = ''.join(_txt(p) for cell in row.cells for p in cell.paragraphs)
        if 'Razem' in row_full: break
        if len(row.cells) > 4 and row.cells[4].paragraphs:
            price_s = _txt(row.cells[4].paragraphs[0]).strip().replace(' ','').replace(' ','')
            try:
                price_f = float(price_s.replace(',','.'))
                _set(row.cells[4].paragraphs[0], f'{price_f:.2f}'.replace('.',','))
            except: pass
    log('  ✓ Цены: 2 знака после запятой')

    # ── Vertical center alignment for all item table cells ─────────────────────
    from docx.enum.table import WD_ALIGN_VERTICAL
    for row in doc.tables[0].rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    log('  ✓ Выравнивание по центру в таблице')

    doc.save(dst)
    log(f'  💾 DOCX сохранён: {os.path.basename(dst)}')

# ─── XLSX HELPERS ─────────────────────────────────────────────────────────────

def _col_n2l(n):
    r = ''
    while n > 0:
        n, rem = divmod(n-1, 26); r = chr(65+rem)+r
    return r

def _col_l2n(s):
    n = 0
    for c in s.upper(): n = n*26+(ord(c)-64)
    return n

def _upd_formula(formula, del_col=3, del_rows_from=2, del_rows_count=2):
    """Adjust formula cell references after deleting rows/column."""
    def repl(m):
        ca, cs, ra, rs = m.groups()
        cn = _col_l2n(cs); rn = int(rs)
        if cn > del_col: cn -= 1
        if rn >= del_rows_from + del_rows_count: rn -= del_rows_count
        return ca + _col_n2l(cn) + ra + str(rn)
    return re.sub(r'(\$?)([A-Z]+)(\$?)(\d+)', repl, formula)

NEW_HEADERS = {
    1:'№', 2:'Article', 3:' Description', 4:'Producer', 5:'COO',
    6:'HS code', 7:'Quantity packs', 8:'Quantity\npcs',
    9:'Net weight, kg', 10:'Gross weight, kg', 11:' E-Price', 12:'Total EUR',
    14:'Pieces per pal', 15:'Quantity pl',
    17:'Цена за 1шт, Евро / E-Price', 18:'Итого, Евро/ Total EUR',
}

# Column widths matching the target layout (landscape A4)
COL_WIDTHS = {
    'A': 2.664, 'B': 10.0,  'C': 26.441, 'D': 13.441, 'E': 7.109,
    'F': 7.332, 'G': 6.109, 'H': 7.555,  'I': 6.887,  'J': 6.555,
    'K': 7.664, 'M': 0.441, 'N': 7.332,  'O': 8.0,    'P': 9.109,
    'R': 6.664, 'S': 7.332,
}

# Russian packaging words → English, with (stem, singular, plural).
# Ordered so carton forms are consumed before the shorter "transport box" stem.
_PLACE_WORDS = [
    (r'обрешет\w*',  'crate',         'crates'),
    (r'паллет\w*',   'pallet',        'pallets'),
    (r'палл\w*',     'pallet',        'pallets'),
    (r'поддон\w*',   'pallet',        'pallets'),
    (r'коробк\w*',   'carton',        'cartons'),
    (r'коробок',     'carton',        'cartons'),
    (r'кор\.',       'carton',        'cartons'),
    (r'короб\w*',    'transport box', 'transport boxes'),
    (r'мест\w*',     'unit',          'units'),
]

# Cyrillic letters that are visual twins of Latin ones (data-entry artefacts in
# product names / dimensions, e.g. "НТMM" → "HTMM", "58х250" → "58x250").
_HOMOGLYPHS = str.maketrans({
    'А':'A','В':'B','Е':'E','К':'K','М':'M','Н':'H','О':'O','Р':'P',
    'С':'C','Т':'T','Х':'X','а':'a','е':'e','к':'k','м':'m','о':'o',
    'р':'p','с':'c','х':'x','у':'y',
})

def _delatinize(text):
    """Convert stray Cyrillic homoglyphs to their Latin twins."""
    return text.translate(_HOMOGLYPHS)

def _translate_places(text):
    """Replace Russian packaging words with English, matching singular/plural to
    the number that precedes the word ("1 поддон" → "1 pallet",
    "4 поддона" → "4 pallets"). Unnumbered occurrences default to plural."""
    for stem, sing, plur in _PLACE_WORDS:
        def _num(m, s=sing, p=plur):
            n = int(m.group(1))
            return f'{m.group(1)}{m.group(2)}{s if n == 1 else p}'
        text = re.sub(r'(\d+)(\s*)' + stem, _num, text, flags=re.IGNORECASE)
        text = re.sub(stem, plur, text, flags=re.IGNORECASE)
    text = re.sub(r'\bчаст\w*', 'part of', text, flags=re.IGNORECASE)
    return text

# ─── XLSX MAIN ────────────────────────────────────────────────────────────────

def process_xlsx(src, dst, params, log):
    import openpyxl
    from openpyxl.worksheet.pagebreak import Break

    spec_num  = params['spec_num']
    date_str  = params['date']
    inv_num   = params['invoice_num']

    log('📊 Обрабатываю XLSX...')
    wb = openpyxl.load_workbook(src, data_only=False)
    ws = wb.active

    # ── Detect header row & "Наименование" column dynamically ────────────────
    hdr_row = None
    name_col = 3
    for r in range(2, 13):
        joined = ' '.join(str(ws.cell(r, c).value) for c in range(1, 20)
                          if ws.cell(r, c).value not in (None, ''))
        if (('Артикул' in joined or 'Artikel' in joined or 'Item' in joined)
                and 'Наименование' in joined):
            hdr_row = r
            for c in range(1, 20):
                v = ws.cell(r, c).value
                if isinstance(v, str) and v.strip() == 'Наименование':
                    name_col = c
                    break
            break
    if hdr_row is None:
        hdr_row = 5

    # Target layout: row1 title, row2 blank, row3 header, row4+ data.
    DEL_COL = name_col
    DEL_FROM = 2
    DEL_CNT = max(0, hdr_row - 3)

    def new_row(r):
        if r < DEL_FROM: return r
        if r < DEL_FROM+DEL_CNT: return None
        return r - DEL_CNT

    def new_col(c):
        if c == DEL_COL: return None
        return c-1 if c > DEL_COL else c

    # Pre-compute corrected formulas keyed by NEW (row, col)
    formula_fixes = {}
    for r in range(1, ws.max_row+1):
        nr = new_row(r)
        if nr is None: continue
        for c in range(1, ws.max_column+1):
            nc = new_col(c)
            if nc is None: continue
            v = ws.cell(r, c).value
            if isinstance(v, str) and v.startswith('='):
                formula_fixes[(nr, nc)] = '=' + _upd_formula(v[1:], DEL_COL, DEL_FROM, DEL_CNT)

    # Structural deletions
    ws.delete_rows(DEL_FROM, DEL_CNT)
    ws.delete_cols(DEL_COL, 1)

    # Apply formula corrections (openpyxl doesn't auto-update refs)
    for (r, c), f in formula_fixes.items():
        ws.cell(r, c).value = f

    # ── Update title (row 1, col B) ───────────────────────────────────────────
    ws.cell(1, 2).value = (
        f'Specification No {spec_num} dated {date_str} to invoice {inv_num}'
    )
    log(f'  ✓ Заголовок: Spec No {spec_num}, {date_str}')

    # ── Update headers (row 3 = was row 5) ────────────────────────────────────
    for c, h in NEW_HEADERS.items():
        ws.cell(3, c).value = h
    log('  ✓ Заголовки переведены на английский')

    # ── Replace «Германия» → «DE» in country column (col 5 after del) ─────────
    cnt = 0
    from openpyxl.styles import Alignment
    for r in range(4, ws.max_row+1):
        if ws.cell(r, 5).value == 'Германия':
            cell = ws.cell(r, 5)
            cell.value = 'DE'
            cell.alignment = Alignment(horizontal='center', vertical='center')
            cnt += 1
    log(f'  ✓ «Германия» → «DE» ({cnt} ячеек, выравнивание по центру)')

    # ── Translate unit cells: кг→kg, шт→pcs (col 5 = was col 6 before del) ───
    for r in range(1, ws.max_row+1):
        v = ws.cell(r, 5).value
        if v == 'кг':
            ws.cell(r, 5).value = 'kg'
        elif v == 'шт':
            ws.cell(r, 5).value = 'pcs'

    # ── Update footer Russian labels + translate Russian packaging words ───────
    for r in range(1, ws.max_row+1):
        for c in range(1, 16):
            v = ws.cell(r, c).value
            if not isinstance(v, str): continue

            # Delivery / payment / shipment terms (col B)
            if 'Условия поставки' in v:
                ws.cell(r, c).value = 'Delivery terms - FCA Warszawa '
                if r > 1 and ws.cell(r-1, c).value in (None, ''):
                    ws.cell(r-1, c).value = f'Shipment date: {date_str}'
                    log(f'  ✓ Дата отгрузки: {date_str}')
            elif 'Условия платежа' in v:
                ws.cell(r, c).value = 'Payment terms: 100% prepayment'
            elif 'Дата отгрузки' in v:
                nv = v.replace('Дата отгрузки', 'Shipment date')
                if date_str:
                    nv = re.sub(r'\d{2}\.\d{2}\.\d{4}', date_str, nv)
                ws.cell(r, c).value = nv
            elif v.strip() in ('Итого:', 'ИТОГО'):
                ws.cell(r, c).value = None

            # Translate Russian packaging words in any cell text
            translated = _translate_places(ws.cell(r, c).value or v)
            if translated != v:
                ws.cell(r, c).value = translated

    # ── Normalize stray Cyrillic homoglyphs in the Description column (col 3) ──
    for r in range(4, ws.max_row+1):
        lp = ws.cell(r, 1).value
        if lp is not None and str(lp).replace('.', '').isdigit():
            d = ws.cell(r, 3).value
            if isinstance(d, str):
                ws.cell(r, 3).value = _delatinize(d)

    log('  ✓ Футер переведён (kg/pcs, паллет/поддон/короб/обрешет/места→EN)')

    # ── Add Seller block at bottom ─────────────────────────────────────────────
    last = ws.max_row
    ws.cell(last + 2, 10).value = 'Seller'
    ws.cell(last + 2, 12).value = 'AOW GROUP SP. z o.o.'
    ws.cell(last + 3, 12).value = 'NIP/TIN  PL5272934015'
    log('  ✓ Добавлен блок Seller')

    # ── Page setup: landscape A4 ───────────────────────────────────────────────
    ws.page_setup.orientation = 'landscape'
    ws.page_setup.paperSize   = 9   # A4
    log('  ✓ Ориентация: альбомная (A4)')

    # ── Print area: columns A–O only ──────────────────────────────────────────
    ws.print_area = f'$A$1:$O${ws.max_row}'

    # ── Page break after column O (15 = Quantity pl) ──────────────────────────
    ws.col_breaks.brk = []   # clear existing
    ws.col_breaks.brk.append(Break(id=15))
    log('  ✓ Разрыв страницы после столбца O (Quantity pl)')

    # ── Column widths ─────────────────────────────────────────────────────────
    for col_letter, width in COL_WIDTHS.items():
        ws.column_dimensions[col_letter].width = width
    log('  ✓ Ширины столбцов установлены')

    # ── Fix number format: 2 decimal places for E-Price and Total EUR ──────────
    for r in range(4, ws.max_row + 1):
        lp = ws.cell(r, 1).value
        if lp and str(lp).replace('.','').isdigit():
            ws.cell(r, 11).number_format = '0.00'   # E-Price
            ws.cell(r, 12).number_format = '0.00'   # Total EUR
    log('  ✓ Формат чисел: 2 знака после запятой')

    # ── Uniform row heights for item data rows only (not footer) ─────────────
    for r in range(4, ws.max_row + 1):
        lp = ws.cell(r, 1).value
        if lp and str(lp).replace('.','').isdigit():
            ws.row_dimensions[r].height = 19.2
    log('  ✓ Высота строк данных выровнена')

    wb.save(dst)
    log(f'  💾 XLSX сохранён: {os.path.basename(dst)}')

def compare_totals(docx_path, xlsx_path, log):
    """Compare final totals between generated DOCX and XLSX."""
    import openpyxl

    # ── DOCX total (from Razem row) ──────────────────────────────────────────
    try:
        from docx import Document
        doc = Document(docx_path)
        docx_total = None
        docx_qty   = None
        for row in doc.tables[0].rows:
            row_full = ''.join(''.join(r.text for r in p.runs)
                               for cell in row.cells for p in cell.paragraphs)
            if 'Razem' in row_full:
                docx_qty   = ''.join(r.text for r in row.cells[3].paragraphs[0].runs).strip()
                docx_total = ''.join(r.text for r in row.cells[5].paragraphs[0].runs).strip()
                break
        if docx_total is None:
            return
        docx_eur = Decimal(docx_total.replace(' ','').replace(' ','').replace(',','.'))
        docx_q   = int(docx_qty.replace(' ','').replace(' ',''))
    except Exception as e:
        log(f'  ⚠ Не удалось прочитать итог DOCX: {e}', 'warn')
        return

    # ── XLSX total (price × qty for each row) ────────────────────────────────
    try:
        wb = openpyxl.load_workbook(xlsx_path, data_only=True)
        ws = wb.active
        xlsx_total = Decimal('0')
        xlsx_qty   = 0
        for r in range(4, ws.max_row + 1):
            lp    = ws.cell(r, 1).value
            qty   = ws.cell(r, 8).value   # Quantity pcs
            price = ws.cell(r, 11).value  # E-Price
            if lp and str(lp).replace('.','').isdigit() and qty and price:
                try:
                    xlsx_total += Decimal(str(qty)) * Decimal(str(price))
                    xlsx_qty   += int(qty)
                except: pass
    except Exception as e:
        log(f'  ⚠ Не удалось прочитать итог XLSX: {e}', 'warn')
        return

    # ── Comparison ───────────────────────────────────────────────────────────
    diff = xlsx_total - docx_eur
    qty_diff = xlsx_qty - docx_q

    log('', None)
    log('─── СВЕРКА ИТОГОВ ───────────────────────────────────', 'hd')
    log(f'  Фактура (DOCX):  {docx_q:>7} шт.   {float(docx_eur):>12,.2f} EUR')
    log(f'  Спецификация:    {xlsx_qty:>7} шт.   {float(xlsx_total):>12,.2f} EUR')

    if abs(diff) < Decimal('0.05') and qty_diff == 0:
        log('  ✅ Суммы совпадают', 'ok')
    else:
        log(f'  ⚠ РАСХОЖДЕНИЕ:', 'warn')
        if qty_diff != 0:
            log(f'     Кол-во: {qty_diff:+d} шт.', 'warn')
        if abs(diff) >= Decimal('0.05'):
            log(f'     Сумма:  {float(diff):+.2f} EUR', 'warn')
    log('─────────────────────────────────────────────────────', 'hd')


# ─── SETTINGS ─────────────────────────────────────────────────────────────────

SETTINGS_FILE = Path.home() / '.aow_converter.json'

def load_cfg():
    try:
        if SETTINGS_FILE.exists():
            return json.loads(SETTINGS_FILE.read_text('utf-8'))
    except: pass
    return {}

def save_cfg(d):
    try: SETTINGS_FILE.write_text(json.dumps(d, ensure_ascii=False, indent=2), 'utf-8')
    except: pass

# ─── GUI ──────────────────────────────────────────────────────────────────────
